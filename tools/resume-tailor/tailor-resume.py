#!/usr/bin/env python3
"""tailor-resume.py — generate a tailored resume .docx from master experience
data + a job-posting YAML.

Reads role_category from the job YAML to drive role-aware scoring.
Falls back to --template for backward compatibility if role_category is absent.

Usage:
    python tools/resume-tailor/tailor-resume.py \\
        --user jane-demo \\
        --job-file applications/jobs/{company}-{job-slug}.yaml \\
        --output-dir applications/jane-demo/{company}-{job-slug}/

    # Scoring only (JSON to stdout, no .docx written):
    python tools/resume-tailor/tailor-resume.py \\
        --user jane-demo \\
        --job-file applications/jobs/{company}-{job-slug}.yaml \\
        --output-dir applications/jane-demo/{company}-{job-slug}/ \\
        --scoring-only
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

import yaml

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")


SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent.parent
DEFAULT_DATA_DIR = REPO_ROOT / "applications" / "resumes" / "data"
DEFAULT_SOURCE_DIR = REPO_ROOT / "applications" / "resumes" / "source"

# Set by CLI args.
DATA_DIR: Path = DEFAULT_DATA_DIR
SOURCE_DIR: Path = DEFAULT_SOURCE_DIR

# Format config — per-user override falls back to global.
def _load_format_config(data_dir: Path) -> dict:
    per_user = data_dir / "format-config.yaml"
    if per_user.exists():
        with open(per_user, encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    global_cfg = SCRIPT_DIR / "format-config.yaml"
    if global_cfg.exists():
        with open(global_cfg, encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    return {}


# Map role_category to the variant key used in summary_variants and title_variants.
ROLE_CATEGORY_TO_VARIANT = {
    "PM": "pm",
    "Product": "product",
    "AI_DT": "ai_dt",
    "AI-DT": "ai_dt",
    "Operations": "operations",
}

# Map role_category to the _focused variant suffix used in achievements.
ROLE_CATEGORY_TO_FOCUSED = {
    "PM": "pm_focused",
    "Product": "product_focused",
    "AI_DT": "ai_dt_focused",
    "AI-DT": "ai_dt_focused",
    "Operations": "operations_focused",
}


def load_yaml(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_master_experience() -> dict:
    return load_yaml(DATA_DIR / "master-experience.yaml")


def load_skills_inventory() -> dict:
    path = DATA_DIR / "skills-inventory.yaml"
    if not path.exists():
        return {"technical_skills": {}}
    return load_yaml(path)


def load_role_schemas() -> dict:
    """Load role-type-schemas.yaml. Per-user version takes priority; fall back
    to the shared file at applications/resumes/data/role-type-schemas.yaml."""
    per_user = DATA_DIR / "role-type-schemas.yaml"
    if per_user.exists():
        return load_yaml(per_user)
    shared = DEFAULT_DATA_DIR / "role-type-schemas.yaml"
    if shared.exists():
        return load_yaml(shared)
    return {"role_categories": {}}


def score_achievement(achievement: dict, requirements: list[str], nice_to_haves: list[str],
                      industry_keywords: list[str], role_category: str | None,
                      schema: dict | None) -> dict:
    """Score an achievement 0-10 against the JD."""
    score = 0
    text_lower = achievement.get("text", "").lower()
    tags = [t.lower() for t in achievement.get("tags", [])]
    tags_joined = " ".join(tags)
    matched_requirements: list[str] = []
    matched_keywords: list[str] = []

    must_have_matches = 0
    for req in requirements:
        req_words = req.lower().split()
        if any(re.search(r"\b" + re.escape(w) + r"\b", text_lower) for w in req_words if len(w) > 2):
            must_have_matches += 1
            matched_requirements.append(req)
        elif any(re.search(r"\b" + re.escape(w) + r"\b", tags_joined) for w in req_words if len(w) > 2):
            must_have_matches += 1
            matched_requirements.append(req)
    score += min(must_have_matches * 3, 9)

    nth_matches = 0
    for nth in nice_to_haves:
        nth_words = nth.lower().split()
        if any(re.search(r"\b" + re.escape(w) + r"\b", text_lower) for w in nth_words if len(w) > 2):
            nth_matches += 1
            matched_requirements.append(nth)
    score += min(nth_matches * 2, 4)

    ind_matches = 0
    for kw in industry_keywords:
        if re.search(r"\b" + re.escape(kw.lower()) + r"\b", text_lower) or kw.lower() in tags_joined:
            ind_matches += 1
            matched_keywords.append(kw)
    score += min(ind_matches, 2)

    if role_category and schema:
        if role_category in achievement.get("better_for", []):
            score += 3
        if role_category in achievement.get("avoid_for", []):
            score -= 5
        if achievement.get("id") in schema.get("achievement_ids_prefer", []):
            score += 2
        if achievement.get("id") in schema.get("achievement_ids_deprioritize", []):
            score -= 3
        boost_matches = 0
        for kw in schema.get("keywords_boost", []):
            if re.search(r"\b" + re.escape(kw.lower()) + r"\b", text_lower):
                matched_keywords.append(kw)
                boost_matches += 1
        score += min(boost_matches, 2)

    base_priority = achievement.get("priority", 5)
    score += min(base_priority // 5, 2)

    return {
        "score": max(min(score, 10), 0),
        "matched_requirements": matched_requirements,
        "matched_keywords": matched_keywords,
    }


def select_achievements(experience: list[dict], requirements: list[str], nice_to_haves: list[str],
                        industry_keywords: list[str], template: str,
                        role_category: str | None = None, schema: dict | None = None,
                        max_per_role: int = 4, force_ids: set | None = None,
                        override_only: bool = False, min_bullets: int | None = None) -> list[dict]:
    """Select and score top achievements per experience entry."""
    selected = []
    focused_variant = ROLE_CATEGORY_TO_FOCUSED.get(role_category, "")
    if force_ids is None:
        force_ids = set()

    for entry in experience:
        achievements = entry.get("achievements", [])
        if not achievements:
            selected.append({**entry, "selected_achievements": [], "scores": []})
            continue

        variant_achievements: list[dict] = []
        seen_base_ids: set = set()

        if focused_variant:
            for ach in achievements:
                if focused_variant in ach.get("variants", []):
                    variant_achievements.append(ach)
                    suffix = ach["id"].rsplit("-", 1)[-1] if "-" in ach["id"] else None
                    if suffix in ("pm", "product", "budget", "stakeholder"):
                        seen_base_ids.add(ach["id"].rsplit("-", 1)[0])

        for ach in achievements:
            if ach in variant_achievements:
                continue
            if ach["id"] in seen_base_ids:
                continue
            variants = ach.get("variants", [])
            if template in variants or "all" in variants or not variants:
                variant_achievements.append(ach)

        for ach in achievements:
            if ach["id"] in force_ids and ach not in variant_achievements:
                variant_achievements.append(ach)

        scored = []
        for ach in variant_achievements:
            result = score_achievement(ach, requirements, nice_to_haves, industry_keywords,
                                       role_category, schema)
            scored.append((result["score"], result, ach))

        scored.sort(key=lambda x: (-x[0], -x[2].get("priority", 0)))

        forced = [(s, r, a) for s, r, a in scored if a["id"] in force_ids]
        non_forced = [(s, r, a) for s, r, a in scored if a["id"] not in force_ids]
        if override_only and force_ids:
            top = forced
        else:
            remaining_slots = max(0, max_per_role - len(forced))
            top = forced + non_forced[:remaining_slots]

        selected.append({
            **entry,
            "selected_achievements": [a for _, _, a in top],
            "scores": [s for s, _, _ in top],
            "match_details": [r for _, r, _ in top],
            "_non_forced_ranked": non_forced,
        })

    # Pass 1: every entry with achievements gets at least 1 bullet.
    for entry in selected:
        has_achievements = bool(entry.get("achievements", []))
        has_selected = bool(entry.get("selected_achievements", []))
        is_career_break = "Career Break" in entry.get("company", "")
        if has_achievements and not has_selected and not is_career_break:
            non_forced = entry.get("_non_forced_ranked", [])
            existing_ids = {a["id"] for a in entry["selected_achievements"]}
            for s, r, a in non_forced:
                if a["id"] not in existing_ids:
                    entry["selected_achievements"].append(a)
                    entry["scores"].append(s)
                    entry["match_details"].append(r)
                    break

    # Pass 2: fill to min_bullets target, distributing across entries.
    fmt = _load_format_config(DATA_DIR)
    target_min = min_bullets if min_bullets is not None else fmt.get("thresholds", {}).get("min_bullets", 22)
    total_bullets = sum(len(e.get("selected_achievements", [])) for e in selected)
    if total_bullets < target_min:
        deficit = target_min - total_bullets
        fill_order = sorted(range(len(selected)), key=lambda i: len(selected[i].get("selected_achievements", [])))
        for _ in range(max_per_role):
            if deficit <= 0:
                break
            for idx in fill_order:
                if deficit <= 0:
                    break
                entry = selected[idx]
                non_forced = entry.get("_non_forced_ranked", [])
                existing_ids = {a["id"] for a in entry["selected_achievements"]}
                if len(entry["selected_achievements"]) >= max_per_role:
                    continue
                for s, r, a in non_forced:
                    if a["id"] not in existing_ids:
                        entry["selected_achievements"].append(a)
                        entry["scores"].append(s)
                        entry["match_details"].append(r)
                        existing_ids.add(a["id"])
                        deficit -= 1
                        break

    for entry in selected:
        entry.pop("_non_forced_ranked", None)

    return selected


def select_summary(master: dict, template: str, role_category: str | None = None) -> str:
    variants = master.get("summary_variants", {})
    variant_key = ROLE_CATEGORY_TO_VARIANT.get(role_category, "")
    if variant_key and variant_key in variants:
        return variants[variant_key].get("text", "")
    if template in variants:
        return variants[template].get("text", "")
    return variants.get("generic", {}).get("text", "")


def select_title(master: dict, template: str, role_category: str | None = None) -> str:
    variants = master.get("summary_variants", {})
    variant_key = ROLE_CATEGORY_TO_VARIANT.get(role_category, "")
    if variant_key and variant_key in variants:
        return variants[variant_key].get("title", "")
    if template in variants:
        return variants[template].get("title", "")
    return variants.get("generic", {}).get("title", "")


def build_skills_section(skills_inventory: dict, role_category: str | None,
                         schema: dict | None) -> str:
    """Build the technical-skills line from skills_inventory.

    Returns an empty 'Technical Skills:' line if no schema or matching skills
    are configured. NEVER returns a hardcoded user-specific tool list.
    """
    if not schema or not role_category:
        return "Technical Skills:"

    allowed_relevance = set(schema.get("skills_filter", []))
    selected = []

    for category_skills in skills_inventory.get("technical_skills", {}).values():
        if not isinstance(category_skills, list):
            continue
        for skill in category_skills:
            skill_relevance = set(skill.get("role_relevance", []))
            if skill_relevance & allowed_relevance:
                selected.append((skill.get("priority", 0), skill.get("name", "")))

    selected.sort(key=lambda x: -x[0])
    skill_names = [name for _, name in selected[:12] if name]

    if not skill_names:
        return "Technical Skills:"

    return "Technical Skills: " + ", ".join(skill_names) + "."


def _tight_para(doc, fmt, space_before=0, space_after=0):
    spacing = fmt.get("spacing", {})
    line_spacing = spacing.get("line_spacing", 13)
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(line_spacing)
    return para


def _tight_bullet(doc, fmt, text, font_size=None):
    spacing = fmt.get("spacing", {})
    fonts = fmt.get("fonts", {})
    if font_size is None:
        font_size = fonts.get("bullet", 11)
    para = doc.add_paragraph(style="List Bullet")
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(spacing.get("bullet_after", 0))
    pf.line_spacing = Pt(spacing.get("line_spacing", 13))
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    return para


def _sanitize_dashes(text: str) -> str:
    return (text.replace("—", " - ")
                .replace("–", " - ")
                .replace("  ", " "))


def build_tailored_docx(master: dict, skills: dict, job: dict, template: str,
                        output_path: Path, role_category: str | None = None,
                        schema: dict | None = None):
    fmt = _load_format_config(DATA_DIR)
    fonts = fmt.get("fonts", {})
    spacing = fmt.get("spacing", {})
    margins = fmt.get("margins", {})
    thresholds = fmt.get("thresholds", {})

    doc = Document()

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(spacing.get("line_spacing", 13))
    style.font.size = Pt(fonts.get("bullet", 11))

    try:
        lb_style = doc.styles["List Bullet"]
        lb_style.paragraph_format.space_before = Pt(0)
        lb_style.paragraph_format.space_after = Pt(spacing.get("bullet_after", 0))
        lb_style.paragraph_format.line_spacing = Pt(spacing.get("line_spacing", 13))
    except KeyError:
        pass

    for section in doc.sections:
        section.top_margin = Inches(margins.get("top", 0.4))
        section.bottom_margin = Inches(margins.get("bottom", 0.4))
        section.left_margin = Inches(margins.get("left", 0.6))
        section.right_margin = Inches(margins.get("right", 0.6))

    contact = master.get("contact", {})
    requirements = job.get("requirements", [])
    nice_to_haves = job.get("nice_to_haves", [])
    industry_keywords = job.get("industry_keywords", [])

    # Name
    name_para = _tight_para(doc, fmt, space_after=spacing.get("after_name", 1))
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_para.add_run(contact.get("full_name", ""))
    name_run.bold = True
    name_run.font.size = Pt(fonts.get("name", 18))

    # Contact line
    contact_text = " | ".join(filter(None, [
        contact.get("location"), contact.get("phone"),
        contact.get("email"), contact.get("linkedin"),
    ]))
    contact_para = _tight_para(doc, fmt, space_after=spacing.get("after_contact", 1))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(fonts.get("contact", 11))

    nationality = contact.get("nationality", "")
    visa_status = contact.get("visa_status", "")
    if nationality or visa_status:
        details_para = _tight_para(doc, fmt, space_after=3)
        details_parts = []
        if nationality:
            details_parts.append(f"Nationality: {nationality}")
        if visa_status:
            details_parts.append(f"Visa: {visa_status}")
        details_run = details_para.add_run(" | ".join(details_parts))
        details_run.font.size = Pt(fonts.get("contact", 11))

    # Title
    title_override = job.get("title_override", "")
    title = title_override.strip() if title_override else select_title(master, template, role_category)
    title_para = _tight_para(doc, fmt,
                              space_before=spacing.get("before_section", 12),
                              space_after=spacing.get("after_title", 2))
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(fonts.get("title", 14))

    # Summary
    summary_override = job.get("summary_override", "")
    summary = summary_override.strip() if summary_override else select_summary(master, template, role_category)
    summary_para = _tight_para(doc, fmt, space_after=spacing.get("after_summary", 3))
    summary_run = summary_para.add_run(summary.strip())
    summary_run.font.size = Pt(fonts.get("summary", 11))

    # Technical Skills
    skills_override = job.get("skills_override", "")
    tech_skills_text = skills_override.strip() if skills_override else build_skills_section(skills, role_category, schema)
    skills_para = _tight_para(doc, fmt, space_after=spacing.get("after_skills", 4))
    skills_run = skills_para.add_run(tech_skills_text)
    skills_run.font.size = Pt(fonts.get("skills", 11))

    # WORK EXPERIENCE heading
    exp_heading = _tight_para(doc, fmt,
                               space_before=spacing.get("before_section", 12), space_after=2)
    exp_run = exp_heading.add_run("WORK EXPERIENCE")
    exp_run.bold = True
    exp_run.font.size = Pt(fonts.get("heading", 14))

    text_overrides = job.get("text_overrides", {}) or {}
    force_ids = set(text_overrides.keys()) if text_overrides else set()
    override_only = bool(job.get("override_only", False))
    job_min_bullets = job.get("min_bullets", None)
    job_max_per_role = job.get("max_per_role", thresholds.get("max_per_role", 4))
    selected = select_achievements(
        master.get("experience", []),
        requirements, nice_to_haves, industry_keywords, template,
        role_category=role_category, schema=schema, force_ids=force_ids,
        override_only=override_only, min_bullets=job_min_bullets,
        max_per_role=job_max_per_role,
    )

    variant_key = ROLE_CATEGORY_TO_VARIANT.get(role_category, template)

    for entry in selected:
        achievements = entry.get("selected_achievements", [])
        title_variants = entry.get("title_variants", {})
        role_title = title_variants.get(variant_key, title_variants.get(template, title_variants.get("all", "")))
        company = entry.get("company", "")
        location = entry.get("location", "")
        dates = entry.get("dates", "")

        if company == "Career Break":
            break_para = _tight_para(doc, fmt, space_before=4, space_after=0)
            break_run = break_para.add_run("CAREER BREAK")
            break_run.bold = True
            break_run.font.size = Pt(fonts.get("company", 11))
            break_para.add_run(f"    {dates}").font.size = Pt(fonts.get("role", 10))
            desc_para = _tight_para(doc, fmt, space_after=2)
            desc_para.add_run(entry.get("description", "")).font.size = Pt(fonts.get("description", 10))
            continue

        if "(continued" not in company:
            company_para = _tight_para(doc, fmt, space_before=4, space_after=0)
            company_run = company_para.add_run(f"{company.upper()} ")
            company_run.bold = True
            company_run.font.size = Pt(fonts.get("company", 11))
            role_run = company_para.add_run(f"{role_title}, {location}")
            role_run.italic = True
            role_run.font.size = Pt(fonts.get("role", 10))
            company_para.add_run(f"    {dates}").font.size = Pt(fonts.get("role", 10))

            desc = entry.get("description", "")
            desc_variants = entry.get("description_variants", {})
            if desc_variants:
                desc = desc_variants.get(variant_key, desc_variants.get(template, desc_variants.get("generic", desc)))
            if desc:
                desc = _sanitize_dashes(desc)
                desc_para = _tight_para(doc, fmt, space_after=0)
                desc_run = desc_para.add_run(desc)
                desc_run.italic = True
                desc_run.font.size = Pt(fonts.get("description", 10))

        for ach in achievements:
            ach_text = text_overrides.get(ach.get("id", ""), ach["text"])
            ach_text = _sanitize_dashes(ach_text)
            _tight_bullet(doc, fmt, ach_text)

    # EDUCATION
    edu_heading = _tight_para(doc, fmt,
                               space_before=spacing.get("before_section", 12), space_after=2)
    edu_run = edu_heading.add_run("EDUCATION")
    edu_run.bold = True
    edu_run.font.size = Pt(fonts.get("heading", 14))

    for edu in master.get("education", []):
        edu_para = _tight_para(doc, fmt, space_after=spacing.get("after_edu_entry", 1))
        edu_run = edu_para.add_run(f"{edu.get('degree', '')} ")
        edu_run.bold = True
        edu_run.font.size = Pt(fonts.get("education", 11))
        edu_para.add_run(
            f"{edu.get('institution', '')} | {edu.get('location', '')} | {edu.get('date', '')}"
        ).font.size = Pt(fonts.get("education", 11))

    # CERTIFICATIONS
    cert_heading = _tight_para(doc, fmt,
                                space_before=spacing.get("before_section", 12), space_after=2)
    cert_run = cert_heading.add_run("CERTIFICATIONS")
    cert_run.bold = True
    cert_run.font.size = Pt(fonts.get("heading", 14))

    for cert in master.get("certifications", []):
        variants = cert.get("variants", [])
        if variant_key in variants or template in variants or not variants:
            cert_para = _tight_para(doc, fmt, space_after=spacing.get("after_cert_entry", 1))
            cert_run = cert_para.add_run(f"{cert.get('name', '')} ")
            cert_run.bold = True
            cert_run.font.size = Pt(fonts.get("certification", 11))
            cert_para.add_run(
                f"{cert.get('issuer', '')} | {cert.get('location', '')} | {cert.get('date', '')}"
            ).font.size = Pt(fonts.get("certification", 11))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Tailored resume saved to: {output_path}")
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Tailor resume for a specific job posting")
    parser.add_argument("--job-file", required=True, help="Path to job posting YAML")
    parser.add_argument("--template",
                        choices=["innovation", "delivery", "generic", "pm", "product", "ai_dt", "operations"],
                        default=None,
                        help="Resume template variant (overridden by role_category in job YAML)")
    parser.add_argument("--output-dir", required=True, help="Output directory for tailored resume")
    parser.add_argument("--scoring-only", action="store_true",
                        help="Run scoring/selection only and output JSON to stdout")
    parser.add_argument("--user", default=os.environ.get("JOB_HUNT_USER"),
                        help="User slug. Falls back to JOB_HUNT_USER env var.")
    parser.add_argument("--data-dir", default=None, help="Override path to per-user data dir")
    parser.add_argument("--source-dir", default=None, help="Override path to source .docx dir")
    args = parser.parse_args()

    global DATA_DIR, SOURCE_DIR
    if args.data_dir:
        DATA_DIR = Path(args.data_dir)
    elif args.user:
        candidate = DEFAULT_DATA_DIR / args.user
        if candidate.exists():
            DATA_DIR = candidate
        else:
            sys.exit(
                f"ERROR: No data dir for user '{args.user}' at {candidate}.\n"
                f"Run /onboard-user or copy applications/resumes/data/_template/ -> {args.user}/."
            )
    else:
        sys.exit(
            "ERROR: --user not provided and JOB_HUNT_USER env var is unset.\n"
            "Pass --user <slug> or set JOB_HUNT_USER in .env."
        )

    if args.source_dir:
        SOURCE_DIR = Path(args.source_dir)
    elif args.user:
        candidate = DEFAULT_SOURCE_DIR / args.user
        if candidate.exists():
            SOURCE_DIR = candidate

    master = load_master_experience()
    skills = load_skills_inventory()
    job = load_yaml(Path(args.job_file))

    role_category = job.get("role_category", None)
    template = args.template or job.get("template", "generic")

    schema = None
    if role_category:
        schemas = load_role_schemas()
        schema_key = role_category.replace("-", "_")
        schema = schemas.get("role_categories", {}).get(schema_key, None)
        if not schema:
            schema = schemas.get("role_categories", {}).get(role_category, None)

    requirements = job.get("requirements", [])
    nice_to_haves = job.get("nice_to_haves", [])
    industry_keywords = job.get("industry_keywords", [])

    text_overrides = job.get("text_overrides", {}) or {}
    force_ids = set(text_overrides.keys()) if text_overrides else set()
    override_only = bool(job.get("override_only", False))
    job_min_bullets = job.get("min_bullets", None)
    job_max_per_role = job.get("max_per_role", 4)
    selected = select_achievements(
        master.get("experience", []),
        requirements, nice_to_haves, industry_keywords, template,
        role_category=role_category, schema=schema, force_ids=force_ids,
        override_only=override_only, min_bullets=job_min_bullets,
        max_per_role=job_max_per_role,
    )

    if args.scoring_only:
        summary_used = job.get("summary_override", "").strip() or select_summary(master, template, role_category)
        title_used = job.get("title_override", "").strip() or select_title(master, template, role_category)
        skills_line = job.get("skills_override", "").strip() or build_skills_section(skills, role_category, schema)

        out = {
            "user": args.user,
            "role_category": role_category,
            "template": template,
            "summary_used": summary_used,
            "title_used": title_used,
            "skills_line": skills_line,
            "selected_achievements": [],
        }
        for entry in selected:
            company = entry.get("company", "")
            for ach, s, md in zip(entry.get("selected_achievements", []),
                                   entry.get("scores", []),
                                   entry.get("match_details", [])):
                out["selected_achievements"].append({
                    "company": company,
                    "id": ach.get("id", ""),
                    "text": ach.get("text", ""),
                    "score": s,
                    "matched_requirements": md.get("matched_requirements", []),
                    "matched_keywords": md.get("matched_keywords", []),
                })
        print(json.dumps(out, indent=2))
        return 0

    output_dir = Path(args.output_dir)
    output_path = output_dir / "resume.docx"
    build_tailored_docx(master, skills, job, template, output_path,
                        role_category=role_category, schema=schema)

    print(f"\nUser: {args.user}")
    print(f"Role category: {role_category or 'None (using template)'}")
    print(f"Template used: {template}")
    if schema:
        print(f"Schema loaded: {schema.get('label', 'unknown')}")
    print(f"Requirements: {len(requirements)}, nice-to-haves: {len(nice_to_haves)}, industry keywords: {len(industry_keywords)}")

    print("\n--- Achievement Selection ---")
    for entry in selected:
        company = entry.get("company", "")
        achievements = entry.get("selected_achievements", [])
        scores = entry.get("scores", [])
        if achievements:
            print(f"\n{company}:")
            for ach, s in zip(achievements, scores):
                print(f"  [{s:2d}] {ach['id']}: {ach['text'][:80]}...")
    return 0


if __name__ == "__main__":
    sys.exit(main())
